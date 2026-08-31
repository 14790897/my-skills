#!/usr/bin/env python3
"""docx-work-report 核心生成库。

输出格式严格对齐官方模板包：
  C:\\git-program\\工作记录填报\\Ai组实习生_工作记录填报_每周五下午打包提交\\
  01_每日工作记录模板.docx / 02_每日新增记录模板.docx / 03_周报模板.docx /
  04_任务看板与导师反馈模板.docx / 05_AI_Agent小团队实习生工作记录台账.xlsx

提供 4 个文档生成函数：
  create_daily_report()    — 日报（7节，对齐 01 模板）
  create_new_record()      — 新增记录（6节，对齐 02 模板）
  create_weekly_report()   — 周报（8节，对齐 03 模板）
  create_kanban()          — 任务看板与导师反馈（5表，对齐 04 模板）
  create_ledger()          — [已弃用] 台账改为 openpyxl 追加到汇总 Excel，
                             不再生成 Word（见 SKILL.md「台账生成」一节）；
                             保留此函数仅为兼容旧周脚本。

排版要点（与官方模板一致）：
  - A4 页面，边距 上1.6 / 下1.5 / 左右1.6 cm
  - 文档标题：仿宋 20pt 加粗 #1F4E79 居中；副标题：仿宋 #666666 居中
  - 节标题：仿宋 14pt 加粗 #1F4E79
  - 数据表：仿宋 9.5pt；边框 0.75pt #B7CBD8；表头底纹 #D9EAF7 加粗 #1F4E79；
    数据行首列 #555555、其余黑色
  - 基本信息表：4 列 [标签, 值, 标签, 值]，标签列底纹 #D9EAF7
  - 规则说明框：底纹 #F7FBFD、边框 #C7DCE8

通过独立的数据脚本调用（如 generate_w26.py），不是直接运行本文件。
所有输出文件保存到 OUTPUT_DIR，命名格式见 SKILL.md。
"""

import os
import sys

# 确保可导入 docx
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx lxml")
    sys.exit(1)


OUTPUT_DIR = os.getcwd()

# ─────────────────────────────────────────────
# 官方模板配色 / 字体常量
# ─────────────────────────────────────────────

COLOR_TITLE = '1F4E79'        # 标题、表头文字深蓝
COLOR_SUBTITLE = '666666'     # 副标题灰
COLOR_HEADER_FILL = 'D9EAF7'  # 表头/标签列底纹浅蓝
COLOR_RULE_FILL = 'F7FBFD'    # 规则说明框底纹
COLOR_BORDER = 'B7CBD8'       # 数据表边框
COLOR_RULE_BORDER = 'C7DCE8'  # 规则说明框边框
COLOR_FIRST_COL = '555555'    # 数据行首列灰

FONT_TABLE = '仿宋'  # 表格、标题字体
FONT_BODY = '宋体'   # 正文默认字体


# ─────────────────────────────────────────────
# 日期格式化工具
# ─────────────────────────────────────────────

def format_date(d):
    """将日期格式化为 YYYY/M/D（不补零），如 2026/5/21、2026/6/3。
    用于表格数据行。支持 str（YYYY-MM-DD）或 date 对象。
    """
    if isinstance(d, str):
        # YYYY-MM-DD → YYYY/M/D
        parts = d.split("-")
        if len(parts) != 3:
            return d
        return f"{parts[0]}/{int(parts[1])}/{int(parts[2])}"
    from datetime import date as _date
    if isinstance(d, _date):
        return f"{d.year}/{d.month}/{d.day}"
    return str(d)


def format_date_cn(d):
    """将日期格式化为 YYYY年M月D日，如 2026年8月28日。
    用于基本信息表（对齐官方模板「____年__月__日」）。
    """
    if isinstance(d, str):
        parts = d.split("-")
        if len(parts) != 3:
            return d
        return f"{parts[0]}年{int(parts[1])}月{int(parts[2])}日"
    from datetime import date as _date
    if isinstance(d, _date):
        return f"{d.year}年{d.month}月{d.day}日"
    return str(d)


# ─────────────────────────────────────────────
# 通用工具函数
# ─────────────────────────────────────────────

def _set_font(run, name=FONT_TABLE, size=None, bold=None, color=None):
    """设置 run 的字体（西文 + eastAsia）、字号、加粗、颜色。"""
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_borders(table, color_hex=COLOR_BORDER, sz='6'):
    """设置表格全部边框为单线，sz 单位为 1/8 pt（6 = 0.75pt）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        borders.append(b)
    tblPr.append(borders)


def set_col_widths(table, widths_cm):
    """按官方模板设置各列宽度（cm）。"""
    table.autofit = False
    for idx, w in enumerate(widths_cm):
        table.columns[idx].width = Cm(w)
        for cell in table.columns[idx].cells:
            cell.width = Cm(w)


def add_rule_box(doc, label, bullets):
    """文档头部的规则说明框：1×1 表格，底纹 F7FBFD、边框 C7DCE8。"""
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, COLOR_RULE_BORDER)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, COLOR_RULE_FILL)
    p0 = cell.paragraphs[0]
    _set_font(p0.add_run(label), bold=True, color=COLOR_TITLE)
    for b in bullets:
        p = cell.add_paragraph()
        _set_font(p.add_run('• '), name='Courier New', size=9.5)
        _set_font(p.add_run(b), size=9.5)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths=None):
    """数据表：表头底纹 D9EAF7 加粗 #1F4E79 居中；数据行首列 #555555。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_borders(table)
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _set_font(r, size=9.5, bold=True, color=COLOR_TITLE)
        set_cell_shading(cell, COLOR_HEADER_FILL)
    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, text in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(text)
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_font(r, size=9.5,
                              color=COLOR_FIRST_COL if ci == 0 else None)
    if widths:
        set_col_widths(table, widths)
    doc.add_paragraph()
    return table


def add_info_table(doc, rows):
    """基本信息表：4 列 [标签, 值, 标签, 值]，标签列底纹 D9EAF7。"""
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_borders(table)
    for ri, row in enumerate(rows):
        l1, v1, l2, v2 = row
        for ci, text in enumerate([l1, v1, l2, v2]):
            cell = table.rows[ri].cells[ci]
            cell.text = str(text)
            for p in cell.paragraphs:
                for r in p.runs:
                    if ci % 2 == 0:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        _set_font(r, size=9.5, bold=True, color=COLOR_TITLE)
                    else:
                        _set_font(r, size=9.5)
        for ci in (0, 2):
            set_cell_shading(table.rows[ri].cells[ci], COLOR_HEADER_FILL)
    set_col_widths(table, [3.0, 5.8, 3.0, 5.8])
    doc.add_paragraph()
    return table


def add_heading(doc, text):
    """节标题：仿宋 14pt 加粗 #1F4E79（对齐官方 Heading 2 样式）。"""
    h = doc.add_heading(text, level=2)
    st = h.style
    st.font.name = FONT_TABLE
    st._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT_TABLE)
    st.font.size = Pt(14)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(COLOR_TITLE)
    for r in h.runs:
        _set_font(r, size=14, bold=True, color=COLOR_TITLE)
    return h


def add_para(doc, text, bold=False):
    """正文段落：宋体 10.5pt（继承 Normal）。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    _set_font(r, name=FONT_BODY, size=10.5, bold=bold)
    return p


def new_doc():
    """新建文档：A4，边距 上1.6 / 下1.5 / 左右1.6 cm，正文宋体 10.5。"""
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21)
    s.page_height = Cm(29.7)
    s.top_margin = Cm(1.6)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(1.6)
    s.right_margin = Cm(1.6)
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT_BODY)
    normal.font.size = Pt(10.5)
    return doc


def title_page(doc, title_text, subtitle=None):
    """文档标题（仿宋 20pt 加粗 #1F4E79 居中）+ 副标题（#666666 居中）。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(title_text), size=20, bold=True, color=COLOR_TITLE)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p2.add_run(subtitle), color=COLOR_SUBTITLE)
    return p


def save(doc, filename):
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"  -> {filename}")
    return path


# ─────────────────────────────────────────────
# 日报生成（7节，对齐 01 模板）
# ─────────────────────────────────────────────

DAILY_RULES = [
    '优先写清“输入—处理—输出—验证”链路，避免只写“学习/调试/开会”。',
    '每项任务至少留下一个可检查产物：代码、PR、文档、实验结果、Prompt、评估数据或复盘结论。',
    '阻塞事项必须写明需要谁、何时、以什么形式支持。',
]


def create_daily_report(date_str, name, project, reviewer, status,
                        goals_rows, tasks_rows, experiments_rows,
                        risks_rows, tomorrow_rows, self_eval_rows,
                        hours='8小时'):
    """
    goals_rows:      [[类别, 内容, 验收标准/证据, 完成情况], ...]
    tasks_rows:      [[序号, Agent模块/任务, 输入与方法, 今日产出, 验证方式/指标, 状态与下一步], ...]
    experiments_rows: [[实验ID, 假设/问题, 配置或变更, 结果, 结论与下一步], ...]（5列）或 [['无']]
    risks_rows:      [[问题/风险, 影响范围, 已尝试方案, 需要支持], ...] 或 [['无']]
    tomorrow_rows:   [[优先级, 计划事项, 预期产出, 依赖/风险], ...]
    self_eval_rows:  [[维度, 自评, 导师反馈], ...]
    """
    doc = new_doc()
    title_page(doc, 'AI Agent 小团队实习生每日工作记录',
               '用于记录当天任务、实验、产出、风险与明日计划')
    add_rule_box(doc, '填写规则', DAILY_RULES)

    add_heading(doc, '一、基本信息')
    add_info_table(doc, [
        ['日期', format_date_cn(date_str), '姓名', name],
        ['所在小组/项目', project, '导师/Reviewer', reviewer],
        ['今日工作时长', hours, '整体状态', status],
    ])

    add_heading(doc, '二、今日目标与实际结果')
    add_table(doc, ['类别', '内容', '验收标准/证据', '完成情况'],
              goals_rows, widths=[2.6, 7.0, 5.0, 3.0])

    add_heading(doc, '三、任务执行记录')
    add_table(doc, ['序号', 'Agent模块/任务', '输入与方法', '今日产出',
                    '验证方式/指标', '状态与下一步'],
              tasks_rows, widths=[1.2, 3.5, 4.2, 4.2, 3.5, 3.3])

    add_heading(doc, '四、实验/调试记录')
    if experiments_rows and experiments_rows[0][0] != '无':
        add_table(doc, ['实验ID', '假设/问题', '配置或变更', '结果', '结论与下一步'],
                  experiments_rows, widths=[2.0, 5.0, 5.0, 3.5, 4.0])
    else:
        add_para(doc, '无')

    add_heading(doc, '五、风险、问题与支持需求')
    if risks_rows and risks_rows[0][0] != '无':
        add_table(doc, ['问题/风险', '影响范围', '已尝试方案', '需要支持'],
                  risks_rows, widths=[5.0, 4.5, 5.0, 5.0])
    else:
        add_para(doc, '无')

    add_heading(doc, '六、明日计划')
    add_table(doc, ['优先级', '计划事项', '预期产出', '依赖/风险'],
              tomorrow_rows, widths=[2.2, 7.0, 5.2, 5.0])

    add_heading(doc, '七、自评与导师反馈')
    add_table(doc, ['维度', '自评', '导师反馈'],
              self_eval_rows, widths=[4.0, 7.2, 7.2])

    return save(doc, f'{name}-{date_str}-日报.docx')


# ─────────────────────────────────────────────
# 新增记录生成（6节，对齐 02 模板）
# ─────────────────────────────────────────────

NEW_RULES = [
    '“新增”包括新知识、新需求、新缺陷、新风险、新Prompt、新工具调用方式、新数据样本、新评估指标等。',
    '每条新增记录都需要判断是否值得进入任务看板、知识库、评估集或周报。',
    '记录重点是可复用性，不是流水账。',
]


def create_new_record(date_str, name, project, module,
                      items_rows, knowledge_rows=None,
                      problems_rows=None, assets_rows=None,
                      weekly_rows=None):
    """
    items_rows:      [[编号, 新增类型, 标题, 来源, 价值/影响, 处理建议], ...]
    knowledge_rows:  [[知识点, 适用场景, 关键步骤/规则, 验证证据, 复用位置], ...]
    problems_rows:   [[问题描述, 复现条件, 严重程度, 临时处理, 负责人/截止时间], ...]
    assets_rows:     [[资产类型, 名称/链接, 用途, 维护人, 下一步], ...]
    weekly_rows:     [[事项, 同步位置, 原因, 证据材料], ...]
    """
    doc = new_doc()
    title_page(doc, 'AI Agent 实习生每日新增记录',
               '用于沉淀新发现、新问题、新知识、新资产和改进机会')
    add_rule_box(doc, '适用范围', NEW_RULES)

    add_heading(doc, '一、基本信息')
    add_info_table(doc, [
        ['日期', format_date_cn(date_str), '记录人', name],
        ['关联项目', project, '关联模块', module],
    ])

    add_heading(doc, '二、新增事项清单')
    add_table(doc, ['编号', '新增类型', '标题', '来源', '价值/影响', '处理建议'],
              items_rows, widths=[1.5, 2.8, 4.5, 3.2, 4.2, 4.0])

    add_heading(doc, '三、新知识/新方法沉淀')
    if knowledge_rows:
        add_table(doc, ['知识点', '适用场景', '关键步骤/规则', '验证证据', '复用位置'],
                  knowledge_rows, widths=[4.0, 4.0, 5.2, 3.5, 3.5])
    else:
        add_para(doc, '无')

    add_heading(doc, '四、新问题/缺陷记录')
    if problems_rows:
        add_table(doc, ['问题描述', '复现条件', '严重程度', '临时处理', '负责人/截止时间'],
                  problems_rows, widths=[5.0, 5.0, 2.5, 4.0, 3.5])
    else:
        add_para(doc, '无')

    add_heading(doc, '五、可复用资产登记')
    if assets_rows:
        add_table(doc, ['资产类型', '名称/链接', '用途', '维护人', '下一步'],
                  assets_rows, widths=[3.0, 5.5, 5.0, 3.0, 3.5])
    else:
        add_para(doc, '无')

    add_heading(doc, '六、是否同步到周报')
    if weekly_rows:
        add_table(doc, ['事项', '同步位置', '原因', '证据材料'],
                  weekly_rows, widths=[5.0, 4.0, 5.0, 5.0])
    else:
        add_para(doc, '无')

    return save(doc, f'{name}-{date_str}-新增.docx')


# ─────────────────────────────────────────────
# 周报生成（8节，对齐 03 模板）
# ─────────────────────────────────────────────

WEEKLY_RULES = [
    '先给结论，再给证据。每个亮点最好配一个链接、指标或截图说明。',
    '风险不只写问题，还要写影响、已采取动作和需要的决策。',
    '下周计划按优先级排列，明确可验收产物。',
]


def create_weekly_report(week_id, week_range, name, reviewer, project, total_hours,
                         summary, conclusion_evidence,
                         outputs_rows, risks_rows,
                         tomorrow_rows, eval_rows,
                         capability_rows=None, growth_rows=None):
    """
    summary:            一句话结论
    conclusion_evidence: 本周最重要证据
    outputs_rows:       [[产出类别, 具体产出, 质量/影响证据, 当前状态], ...]
    capability_rows:    [[能力维度, 本周动作, 指标或观察, 结论], ...] 或 None（写"无"）
    risks_rows:         [[问题/风险, 影响, 已采取动作, 需要团队决策/支持], ...]
    growth_rows:        [[本周学到的关键点, 对应实践, 可复用沉淀], ...] 或 None（写"无"）
    tomorrow_rows:      [[优先级, 计划事项, 预期产物, 验收标准, 依赖], ...]
    eval_rows:          [[评价维度, 评分, 建议], ...]
    """
    doc = new_doc()
    title_page(doc, 'AI Agent 小团队实习生周报',
               '用于向导师和团队汇总本周结果、证据、风险与下周计划')
    add_rule_box(doc, '周报原则', WEEKLY_RULES)

    add_heading(doc, '一、基本信息')
    add_info_table(doc, [
        ['周次', week_id, '周期', week_range],
        ['姓名', name, '导师/Reviewer', reviewer],
        ['项目/小组', project, '本周总工时', f'{total_hours}小时'],
    ])

    add_heading(doc, '二、本周一句话总结')
    add_table(doc, ['一句话结论', '本周最重要证据'],
              [[summary, conclusion_evidence]], widths=[9.0, 9.0])

    add_heading(doc, '三、本周关键产出')
    add_table(doc, ['产出类别', '具体产出', '质量/影响证据', '当前状态'],
              outputs_rows, widths=[3.0, 7.0, 6.0, 3.0])

    add_heading(doc, '四、AI Agent 能力建设进展')
    if capability_rows:
        add_table(doc, ['能力维度', '本周动作', '指标或观察', '结论'],
                  capability_rows, widths=[4.0, 5.0, 5.0, 5.0])
    else:
        add_para(doc, '无')

    add_heading(doc, '五、问题、风险与决策需求')
    add_table(doc, ['问题/风险', '影响', '已采取动作', '需要团队决策/支持'],
              risks_rows, widths=[5.0, 4.0, 5.0, 5.0])

    add_heading(doc, '六、个人成长与方法沉淀')
    if growth_rows:
        add_table(doc, ['本周学到的关键点', '对应实践', '可复用沉淀'],
                  growth_rows, widths=[6.0, 6.0, 6.5])
    else:
        add_para(doc, '无')

    add_heading(doc, '七、下周计划')
    add_table(doc, ['优先级', '计划事项', '预期产物', '验收标准', '依赖'],
              tomorrow_rows, widths=[2.0, 5.5, 4.5, 4.5, 3.0])

    add_heading(doc, '八、导师评价')
    add_table(doc, ['评价维度', '评分', '建议'],
              eval_rows, widths=[4.0, 2.2, 12.0])

    return save(doc, f'{name}-{week_id}-周报.docx')


# ─────────────────────────────────────────────
# 任务看板与导师反馈（5表，对齐 04 模板）
# ─────────────────────────────────────────────

KANBAN_RULES = [
    '每周一更新看板优先级，每日站会只更新状态、阻塞和下一步。',
    '导师反馈尽量写成可执行动作：修改什么、标准是什么、什么时候复查。',
    '阶段评价不只看完成数量，还要看证据质量、复盘能力和可复用沉淀。',
]


def create_kanban(week_id, week_range, name, reviewer,
                  tasks_rows, standup_rows, feedback_rows,
                  eval_rows, retro_rows):
    """
    tasks_rows:    [[T-001, 任务名称, 负责人, P0, Done, 验收标准, 截止日期], ...]
    standup_rows:  [[日期, 昨日完成, 今日计划, 阻塞, 需同步事项], ...]
    feedback_rows: [[日期, 反馈对象, 反馈类型, 具体反馈, 行动项, 复查时间], ...] 或 None
    eval_rows:     [[评价维度, 观察证据, 评分, 改进建议], ...]
    retro_rows:    [[复盘问题, 记录], ...]
    """
    doc = new_doc()
    title_page(doc, 'AI Agent 小团队任务看板与导师反馈模板',
               '用于小团队管理、任务流转、Review与阶段评价')
    add_rule_box(doc, '使用方式', KANBAN_RULES)

    add_heading(doc, '一、任务看板')
    add_table(doc, ['任务ID', '任务名称', '负责人', '优先级', '状态',
                    '验收标准', '截止日期'],
              tasks_rows, widths=[1.57, 3.92, 2.09, 1.74, 2.18, 4.53, 2.18])

    add_heading(doc, '二、每日站会记录')
    add_table(doc, ['日期', '昨日完成', '今日计划', '阻塞', '需同步事项'],
              standup_rows, widths=[2.13, 4.17, 4.17, 3.71, 3.71])

    add_heading(doc, '三、导师反馈记录')
    if feedback_rows:
        add_table(doc, ['日期', '反馈对象', '反馈类型', '具体反馈', '行动项', '复查时间'],
                  feedback_rows, widths=[2.04, 2.31, 2.59, 5.09, 4.17, 2.31])
    else:
        add_para(doc, '暂无数据')

    add_heading(doc, '四、阶段评价表')
    add_table(doc, ['评价维度', '观察证据', '评分', '改进建议'],
              eval_rows, widths=[3.7, 6.48, 1.85, 5.55])

    add_heading(doc, '五、团队复盘')
    add_table(doc, ['复盘问题', '记录'],
              retro_rows, widths=[5.02, 11.87])

    return save(doc, f'{name}-{week_id}-任务看板.docx')


# ─────────────────────────────────────────────
# 台账（已弃用，仅兼容旧脚本）
# 台账实际走 openpyxl 追加到汇总 Excel，见 SKILL.md「台账生成」一节。
# 列头已更新为与 05_台账.xlsx 各 sheet 一致。
# ─────────────────────────────────────────────

def create_ledger(week_id, week_range, name, project,
                  daily_rows, new_items_rows,
                  weekly_summary, kanban_rows):
    """
    daily_rows:      [[日期, 姓名, Agent模块, 任务类型, 今日目标, 实际产出, 证据链接,
                      验证方式/指标, 状态, 工时, 阻塞事项, 明日计划, 导师备注], ...]（13列）
    new_items_rows:  [[日期, 新增类型, 关联模块, 标题, 来源, 价值/影响, 处理建议,
                      是否进看板, 是否进周报, 证据/链接, 备注], ...]（11列）
    weekly_summary:  [[周次, 周期, 姓名, 项目/小组, 一句话总结, 关键产出,
                      指标/证据, 问题与风险, 下周计划, 导师评价], ...]（10列）
    kanban_rows:     [[任务ID, 任务名称, 负责人, 优先级, 状态, Agent模块, 验收标准,
                      开始日期, 截止日期, 阻塞/依赖, 备注], ...]（11列）
    """
    doc = new_doc()
    title_page(doc, 'AI Agent 小团队工作记录台账')

    add_para(doc, f'周次：{week_id}（{week_range}）  |  姓名：{name}  |  项目：{project}', bold=True)
    doc.add_paragraph()

    add_heading(doc, '表一：每日工作记录')
    add_table(doc, ['日期', '姓名', 'Agent模块', '任务类型', '今日目标', '实际产出', '证据链接',
                    '验证方式/指标', '状态', '工时', '阻塞事项', '明日计划', '导师备注'], daily_rows)

    add_heading(doc, '表二：每日新增记录')
    add_table(doc, ['日期', '新增类型', '关联模块', '标题', '来源', '价值/影响', '处理建议',
                    '是否进看板', '是否进周报', '证据/链接', '备注'], new_items_rows)

    add_heading(doc, '表三：周报汇总')
    add_table(doc, ['周次', '周期', '姓名', '项目/小组', '一句话总结', '关键产出',
                    '指标/证据', '问题与风险', '下周计划', '导师评价'], weekly_summary)

    add_heading(doc, '表四：任务看板')
    add_table(doc, ['任务ID', '任务名称', '负责人', '优先级', '状态', 'Agent模块', '验收标准',
                    '开始日期', '截止日期', '阻塞/依赖', '备注'], kanban_rows)

    return save(doc, f'{name}-{week_id}-台账.docx')
